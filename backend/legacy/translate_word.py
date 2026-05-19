# translate_word.py
import os
import gc
from concurrent.futures import ThreadPoolExecutor, as_completed
from docx import Document
from docx.oxml.ns import qn
from langchain_core.tools import ToolException
from translate_text import translate_text, _translation_cache
from pathlib import Path
from dotenv import load_dotenv
load_dotenv()

# 获取当前脚本所在目录的父级目录（即 SIS_Agent 根目录）
SIS_AGENT_ROOT = Path(__file__).parent.parent

# ---------------------- 全局配置 ----------------------
DEFAULT_INPUT_DIR = os.getenv("TRANSLATE_INPUT_DIR") or str(SIS_AGENT_ROOT / "translate" / "input")
DEFAULT_TARGET_LANG = os.getenv("TRANSLATE_TARGET_LANG") or "日语"
DEFAULT_DELAY = os.getenv("TRANSLATE_DELAY") or 1.2
MAX_WORKERS = int(os.getenv("MAX_WORKERS") or 6)


def _paragraph_has_picture(paragraph):
    """判断段落是否包含图片"""
    for run in paragraph.runs:
        if run._element.find(qn('w:drawing')) is not None:
            return True
    return False


def _collect_tasks(doc):
    """
    收集所有需要翻译的文本单元，保持原始顺序。
    返回: [(setter, original_text), ...]
    setter 接受关键字参数 trans_text。
    """
    tasks = []

    # 1. 主体段落
    for para in doc.paragraphs:
        if _paragraph_has_picture(para):
            continue
        text = para.text.strip()
        if not text or len(text) <= 1:
            continue
        # 使用默认参数绑定当前 para
        def setter_para(p=para, trans_text=None):
            if trans_text:
                p.text = trans_text
        tasks.append((setter_para, text))

    # 2. 表格单元格内的段落
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if _paragraph_has_picture(para):
                        continue
                    text = para.text.strip()
                    if not text or len(text) <= 1:
                        continue
                    def setter_cell(p=para, trans_text=None):
                        if trans_text:
                            p.text = trans_text
                    tasks.append((setter_cell, text))

    # 3. 文本框内的段落
    root = doc.part.element
    txbx_contents = root.xpath('.//*[local-name()="txbxContent"]')
    for txbx in txbx_contents:
        for para_elem in txbx.xpath('.//*[local-name()="p"]'):
            t_nodes = para_elem.xpath('.//*[local-name()="t"]')
            original = ''.join(t_node.text for t_node in t_nodes if t_node.text).strip()
            if not original or len(original) <= 1:
                continue
            def setter_textbox(nodes=t_nodes, trans_text=None):
                if trans_text:
                    for node in nodes:
                        node.text = trans_text
            tasks.append((setter_textbox, original))

    return tasks


def _translate_group(tasks_group, target_lang, delay, base_context, group_id):
    """
    顺序翻译一组任务，共享组内上下文。
    返回: [(setter, translated_text), ...]
    """
    ctx = {
        "file_name": base_context["file_name"],
        "slide_num": 0,
        "total_slides": base_context["total_slides"],
        "translated_segments": [],
        "term_requirements": base_context["term_requirements"]
    }
    results = []
    for idx, (setter, original) in enumerate(tasks_group):
        try:
            translated = translate_text(original, target_lang, delay, context=ctx)
            if translated is None:
                translated = original
            # 更新组内上下文
            ctx["translated_segments"].append(f"原文：{original} | 译文：{translated}")
            if len(ctx["translated_segments"]) > 5:
                ctx["translated_segments"].pop(0)
            results.append((setter, translated))
        except Exception as e:
            print(f"组 {group_id} 翻译第 {idx+1} 个单元失败: {e}")
            results.append((setter, original))
    return results


def translate_word_file(file_name: str, workspace_dir: str = DEFAULT_INPUT_DIR,
                        target_lang: str = DEFAULT_TARGET_LANG,
                        delay: float = DEFAULT_DELAY) -> str:
    """翻译 Word 文档（仅 .docx），顺序分组并行翻译"""
    ext = os.path.splitext(file_name)[1].lower()
    if ext == ".doc":
        raise ToolException("暂不支持 .doc 格式，请转换为 .docx")
    if ext != ".docx":
        raise ToolException(f"不支持的文件类型: {ext}，仅支持 .docx")

    input_path = os.path.abspath(os.path.join(workspace_dir, file_name))
    output_path = os.path.abspath(os.path.join(
        workspace_dir,
        f"{os.path.splitext(file_name)[0]}_{target_lang}.docx"
    ))

    if not os.path.exists(input_path):
        raise ToolException(f"文件不存在: {input_path}")

    doc = Document(input_path)
    tasks = _collect_tasks(doc)
    total = len(tasks)
    if total == 0:
        doc.save(output_path)
        return "未找到可翻译的纯文本内容"

    # 均匀分组
    group_size = (total + MAX_WORKERS - 1) // MAX_WORKERS
    groups = [tasks[i:i+group_size] for i in range(0, total, group_size)]

    base_context = {
        "file_name": file_name,
        "slide_num": 0,
        "total_slides": total,
        "translated_segments": [],
        "term_requirements": "TBOX译为TBOX、TSU译为TSU、CAN总线译为CANバス（日语）/CAN Bus（英语），公司名称不翻译"
    }

    print(f"共 {total} 个文本单元，分为 {len(groups)} 组并行翻译...")

    # 并行翻译各组
    group_results = [None] * len(groups)
    with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
        futures = {executor.submit(_translate_group, group, target_lang, delay, base_context, idx): idx
                   for idx, group in enumerate(groups)}
        for future in as_completed(futures):
            idx = futures[future]
            try:
                group_results[idx] = future.result()
                print(f"组 {idx+1}/{len(groups)} 翻译完成")
            except Exception as e:
                print(f"组 {idx+1} 翻译失败: {e}")
                group_results[idx] = [(setter, original) for setter, original in groups[idx]]

    # 合并结果
    all_results = []
    for grp in group_results:
        all_results.extend(grp)

    # 单线程顺序写回（使用关键字参数 trans_text）
    print("写回译文...")
    for idx, (setter, translated) in enumerate(all_results):
        try:
            setter(trans_text=translated)   # 关键修正：使用关键字参数
        except Exception as e:
            print(f"写回第 {idx+1} 个单元失败: {e}")
        if (idx + 1) % 50 == 0:
            print(f"已写回 {idx+1}/{total}")

    doc.save(output_path)
    print(f"✅ 翻译完成: {output_path}")

    # _translation_cache.clear()
    gc.collect()
    return f"Word翻译完成！输出路径: {output_path}"